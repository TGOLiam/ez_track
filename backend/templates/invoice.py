import os
from docx import Document
from datetime import datetime

TEMPLATES_DIR = os.path.dirname(__file__)

_TEMPLATES = {
    "sigla": "EzTrack_Non-VAT_Invoice_Template_Simula_Sigla.docx",
    "unlad": "EzTrack_VAT_Invoice_Template_Unlad.docx",
}

def _find_and_replace(doc, placeholder, value):
    value = str(value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            return
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return

def _find_and_replace_all(doc, placeholder, value):
    value = str(value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(r.text for r in paragraph.runs)
                    if placeholder in full_text:
                        new_text = full_text.replace(placeholder, value)
                        for ri, run in enumerate(paragraph.runs):
                            run.text = new_text if ri == 0 else ''
    for paragraph in doc.paragraphs:
        full_text = ''.join(r.text for r in paragraph.runs)
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, value)
            for ri, run in enumerate(paragraph.runs):
                run.text = new_text if ri == 0 else ''

def _find_table_by_header(doc, header_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if header_text.lower() in cell.text.lower():
                    return table
    return None

def _fill_items_table(table, items, col_map):
    if not table or not items:
        return
    header_rows = 1
    data_rows = len(table.rows) - header_rows
    for idx, item in enumerate(items):
        if idx < data_rows:
            row = table.rows[header_rows + idx]
        else:
            row = table.add_row()
        amt = item["qty"] * item["unit_price"]
        for ci, field in col_map.items():
            if ci < len(row.cells):
                val = {
                    "description": item.get("description", ""),
                    "qty": str(item["qty"]),
                    "unit_price": f"₱{item['unit_price']:,.2f}",
                    "amount": f"₱{amt:,.2f}",
                }.get(field, "")
                if val:
                    row.cells[ci].text = val

def build(items, customer, payment, ctx, tier, discount=0, withholding_tax=0, notes=""):
    tmpl = _TEMPLATES.get(tier, _TEMPLATES["sigla"])
    doc = Document(os.path.join(TEMPLATES_DIR, tmpl))
    biz = ctx.get("business") or {}
    biz_name = biz.get("name", "My Business")
    biz_city = biz.get("city", "")
    contact = biz.get("contact", "09XX-XXX-XXXX")

    _find_and_replace(doc, "[Your Business Name]", biz_name)
    _find_and_replace(doc, "[Business Address", biz_city)
    _find_and_replace(doc, "Barangay, City]", biz_city)
    _find_and_replace(doc, "09XX-XXX-XXXX", contact)
    _find_and_replace(doc, "[Month DD, YYYY]", datetime.now().strftime("%B %d, %Y"))
    _find_and_replace(doc, "[0000001]", str(ctx.get("nextDocNumber", 1)))
    _find_and_replace(doc, "[Customer Name]", customer.get("name", ""))

    items_table = _find_table_by_header(doc, "ITEM")
    if not items_table:
        items_table = _find_table_by_header(doc, "QTY")

    if items_table:
        col_map = {0: "description", 1: "qty", 2: "unit_price", 3: "amount"}
        _fill_items_table(items_table, items, col_map)

    subtotal = sum(i["qty"] * i["unit_price"] for i in items)
    total_due = subtotal - discount - withholding_tax

    _find_and_replace_all(doc, "₱[0.00]", f"₱{total_due:,.2f}")

    if payment.get("method"):
        _find_and_replace(doc, "Cash", payment["method"])
    if payment.get("status") and "paid" in payment["status"].lower():
        _find_and_replace_all(doc, "☐", "✓")

    if tier == "unlad":
        buyer_tin = customer.get("tin", "000-000-000-000")
        buyer_addr = customer.get("address", "")
        _find_and_replace(doc, "[Buyer's Registered Name]", customer.get("name", ""))
        _find_and_replace(doc, "[Buyer's TIN", buyer_tin)
        _find_and_replace(doc, "[Buyer's Address]", buyer_addr)
        vat_sales = subtotal / 1.12 if subtotal else 0
        vat_amt = subtotal - vat_sales
        _find_and_replace_all(doc, "₱[0.00]", f"₱{vat_sales:,.2f}")

    if notes:
        _find_and_replace(doc, "[e.g. Please settle", notes)

    return doc
