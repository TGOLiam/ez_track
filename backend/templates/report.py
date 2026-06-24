import os
import re
from docx import Document
from datetime import datetime

TEMPLATES_DIR = os.path.dirname(__file__)
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "EzTrack_Sales_Report_Template_DevSpec.docx")

def _replace_placeholder(doc, tag, value):
    value = str(value)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if tag in run.text:
                run.text = run.text.replace(tag, value)
                return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if tag in run.text:
                            run.text = run.text.replace(tag, value)
                            return

def _add_table_row(table, cells_data):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        if i < len(row.cells):
            row.cells[i].text = text
    return row

def build(ctx, tier, data):
    doc = Document(TEMPLATE_PATH)
    biz = ctx.get("business") or {}
    user = ctx.get("user") or {}

    now = datetime.now()
    _replace_placeholder(doc, "{{generated_date}}", now.strftime("%B %d, %Y at %I:%M %p"))
    _replace_placeholder(doc, "{{business_name}}", biz.get("name", "My Business"))
    _replace_placeholder(doc, "{{owner_name}}", user.get("name", "User"))
    _replace_placeholder(doc, "{{report_period_label}}", data.get("period", ""))
    _replace_placeholder(doc, "{{report_type}}", "Weekly Report" if tier == "sigla" else "Monthly Report")
    _replace_placeholder(doc, "{{tier_name}}", "Sigla" if tier == "sigla" else "Unlad")

    _replace_placeholder(doc, "{{ai_summary_text}}", data.get("ai_summary", ""))
    _replace_placeholder(doc, "{{total_sales}}", f"₱{data.get('total_sales', 0):,.2f}")
    _replace_placeholder(doc, "{{total_expenses}}", f"₱{data.get('total_expenses', 0):,.2f}")
    net = (data.get("total_sales", 0) or 0) - (data.get("total_expenses", 0) or 0)
    _replace_placeholder(doc, "{{net_remaining}}", f"₱{net:,.2f}")
    _replace_placeholder(doc, "{{comparison_vs_last_period}}", data.get("comparison_text", ""))

    _replace_placeholder(doc, "{{strongest_period_label}}", data.get("strongest_period_label", ""))
    _replace_placeholder(doc, "{{strongest_period_amount}}", data.get("strongest_period_amount", ""))
    _replace_placeholder(doc, "{{weakest_period_label}}", data.get("weakest_period_label", ""))
    _replace_placeholder(doc, "{{weakest_period_amount}}", data.get("weakest_period_amount", ""))

    cats = data.get("categories", [])
    for cat in cats:
        cat_table = None
        for table in doc.tables:
            for row in table.rows:
                if "Category" in row.cells[0].text if row.cells else "":
                    cat_table = table
                    break
            if cat_table:
                break
        if cat_table:
            _add_table_row(cat_table, [
                cat.get("name", ""),
                f"₱{cat.get('amount', 0):,.2f}",
                cat.get("pct", "0%"),
            ])

    _replace_placeholder(doc, "{{tip_1}}", data.get("tip_1", ""))
    _replace_placeholder(doc, "{{tip_2}}", data.get("tip_2", ""))

    if tier == "unlad":
        _replace_placeholder(doc, "{{gross_sales}}", f"₱{data.get('gross_sales', data.get('total_sales', 0)):,.2f}")
        _replace_placeholder(doc, "{{total_direct_costs}}", f"₱{data.get('total_direct_costs', 0):,.2f}")
        gp = (data.get("gross_sales", data.get("total_sales", 0)) or 0) - (data.get("total_direct_costs", 0) or 0)
        _replace_placeholder(doc, "{{gross_profit}}", f"₱{gp:,.2f}")
        _replace_placeholder(doc, "{{total_opex}}", f"₱{data.get('total_opex', data.get('total_expenses', 0)):,.2f}")
        np = gp - (data.get("total_opex", data.get("total_expenses", 0)) or 0)
        _replace_placeholder(doc, "{{net_profit}}", f"₱{np:,.2f}")
        _replace_placeholder(doc, "{{forecast_30day_position}}", data.get("forecast", ""))
        _replace_placeholder(doc, "{{cash_danger_flag}}", data.get("cash_risk", "None"))
        _replace_placeholder(doc, "{{goal_label}}", data.get("goal_label", ""))
        _replace_placeholder(doc, "{{goal_progress_pct}}", data.get("goal_progress", ""))
        _replace_placeholder(doc, "{{goal_status_note}}", data.get("goal_note", ""))
        _replace_placeholder(doc, "{{ar_current}}", data.get("ar_current", "₱0.00"))
        _replace_placeholder(doc, "{{ar_30days}}", data.get("ar_30days", "₱0.00"))
        _replace_placeholder(doc, "{{ar_60days}}", data.get("ar_60days", "₱0.00"))
        _replace_placeholder(doc, "{{ar_90plus_days}}", data.get("ar_90plus", "₱0.00"))

    return doc
