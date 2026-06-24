import os
from docx import Document
from datetime import datetime

TEMPLATES_DIR = os.path.dirname(__file__)

_TEMPLATES = {
    "simula": "EzTrack_Basic_Receipt_Template.docx",
    "sigla": "EzTrack_Official_Receipt_NON-VAT.docx",
    "unlad": "EzTrack_Official_Receipt_VAT.docx",
}

def _find_and_replace(doc, placeholder, value):
    value = str(value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            return
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return

def _find_and_replace_all(doc, placeholder, value):
    value = str(value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(r.text for r in paragraph.runs)
                    if placeholder in full_text:
                        new_text = full_text.replace(placeholder, value)
                        for ri, run in enumerate(paragraph.runs):
                            run.text = new_text if ri == 0 else ''
    for paragraph in doc.paragraphs:
        full_text = ''.join(r.text for r in paragraph.runs)
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, value)
            for ri, run in enumerate(paragraph.runs):
                run.text = new_text if ri == 0 else ''

def _find_table_by_header(doc, header_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if header_text.lower() in cell.text.lower():
                    return table
    return None

def _fill_items_table(table, items, col_map):
    if not table or not items:
        return
    header_rows = 1
    data_rows = len(table.rows) - header_rows
    for idx, item in enumerate(items):
        if idx < data_rows:
            row = table.rows[header_rows + idx]
        else:
            row = table.add_row()
        amt = item["qty"] * item["unit_price"]
        for ci, field in col_map.items():
            if ci < len(row.cells):
                val = {
                    "description": item.get("description", ""),
                    "qty": str(item["qty"]),
                    "unit_price": f"₱{item['unit_price']:,.2f}",
                    "amount": f"₱{amt:,.2f}",
                }.get(field, "")
                if val:
                    row.cells[ci].text = val

def _amount_to_words(amount):
    if amount == 0:
        return "Zero Pesos Only"
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
            "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
            "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    def _convert(n):
        if n < 20:
            return ones[n]
        if n < 100:
            return tens[n // 10] + (" " + ones[n % 10] if n % 10 else "")
        if n < 1000:
            return ones[n // 100] + " Hundred" + (" " + _convert(n % 100) if n % 100 else "")
        if n < 1000000:
            return _convert(n // 1000) + " Thousand" + (" " + _convert(n % 1000) if n % 1000 else "")
        return ""
    whole = int(amount)
    cents = round((amount - whole) * 100)
    result = _convert(whole) + " Pesos"
    if cents:
        result += " and " + _convert(cents) + " Centavos"
    return result + " Only"

def build(items, customer, payment, ctx, tier):
    tmpl = _TEMPLATES.get(tier, _TEMPLATES["simula"])
    doc = Document(os.path.join(TEMPLATES_DIR, tmpl))
    biz = ctx.get("business") or {}
    biz_name = biz.get("name", "My Business")
    biz_city = biz.get("city", "")
    contact = biz.get("contact", "09XX-XXX-XXXX")

    _find_and_replace(doc, "[Your Business Name]", biz_name)
    _find_and_replace(doc, "[Business Address / Barangay, City]", biz_city)
    _find_and_replace(doc, "Barangay, City]", biz_city)
    _find_and_replace(doc, "09XX-XXX-XXXX", contact)
    _find_and_replace(doc, "[Month DD, YYYY]", datetime.now().strftime("%B %d, %Y"))
    _find_and_replace(doc, "[0001]", str(ctx.get("nextDocNumber", 1)))
    _find_and_replace(doc, "[Customer Name]", customer.get("name", ""))
    if customer.get("contact"):
        _find_and_replace(doc, "[09XX-XXX-XXXX]", customer["contact"])

    items_table = _find_table_by_header(doc, "AMOUNT")
    if not items_table:
        items_table = _find_table_by_header(doc, "WHAT WAS PAID")

    if items_table:
        ncols = len(items_table.rows[0].cells)
        col_map = {2: "qty", 3: "unit_price", 4: "amount"} if ncols > 4 else {1: "amount"}
        col_map[0] = "description"
        _fill_items_table(items_table, items, col_map)

    total = sum(i["qty"] * i["unit_price"] for i in items)
    _find_and_replace_all(doc, "₱[0.00]", f"₱{total:,.2f}")

    words = _amount_to_words(total)
    for para in doc.paragraphs:
        full = ''.join(r.text for r in para.runs)
        if "[Amount in words" in full or "One Thousand" in full or "Amount in Words" in full:
            for ri, run in enumerate(para.runs):
                run.text = words if ri == 0 else ''
            break

    if payment.get("method"):
        _find_and_replace(doc, "Cash", payment["method"])
    if payment.get("status") and "paid" in payment["status"].lower():
        _find_and_replace_all(doc, "☐", "✓")

    if tier == "unlad":
        biz_tin = biz.get("tin", "000-000-000-000")
        _find_and_replace(doc, "[000-000-000-000]", biz_tin)

    return doc
